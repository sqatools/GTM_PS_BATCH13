"""
Document Loader
Handles loading documents from various formats (PDF, Word, Text)
"""

import os
from pathlib import Path
from typing import List
from unittest.mock import Mock

# Gracefully handle missing langchain
try:
    from langchain_core.documents import Document
except ModuleNotFoundError:
    # Create mock Document class for when langchain is not installed
    class Document(Mock):
        def __init__(self, page_content: str = "", metadata: dict = None, **kwargs):
            super().__init__()
            self.page_content = page_content
            self.metadata = metadata or {}

from app.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """Base document loader"""
    
    def __init__(self, data_path: str = "app/data/raw_docs"):
        self.data_path = data_path
        logger.info(f"Initialized DocumentLoader with path: {data_path}")
    
    def load_documents(self) -> List[Document]:
        """Load all documents from data path"""
        documents = []
        
        if not os.path.exists(self.data_path):
            logger.warning(f"Data path not found: {self.data_path}")
            return documents
        
        # Iterate through all files in data path
        for filename in os.listdir(self.data_path):
            filepath = os.path.join(self.data_path, filename)
            
            if filename.endswith('.docx'):
                logger.info(f"Loading Word document: {filename}")
                docs = self.load_docx(filepath)
                documents.extend(docs)
            
            elif filename.endswith('.pdf'):
                logger.info(f"Loading PDF: {filename}")
                docs = self.load_pdf(filepath)
                documents.extend(docs)
            
            elif filename.endswith(('.txt', '.md')):
                logger.info(f"Loading text file: {filename}")
                docs = self.load_text(filepath)
                documents.extend(docs)
        
        logger.info(f"Total documents loaded: {len(documents)}")
        return documents
    
    @staticmethod
    def load_docx(filepath: str) -> List[Document]:
        """Load Word (.docx) document"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(filepath)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            # Create Document object
            doc_obj = Document(
                page_content=full_text,
                metadata={
                    "source": filepath,
                    "filename": Path(filepath).name,
                    "type": "docx"
                }
            )
            
            logger.info(f"Successfully loaded Word document: {filepath}")
            return [doc_obj]
        
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return []
        except Exception as e:
            logger.error(f"Error loading Word document {filepath}: {str(e)}")
            return []
    
    @staticmethod
    def load_pdf(filepath: str) -> List[Document]:
        """Load PDF document"""
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(filepath)
            documents = loader.load()
            
            # Add source metadata
            for doc in documents:
                doc.metadata["source"] = filepath
                doc.metadata["filename"] = Path(filepath).name
            
            logger.info(f"Successfully loaded PDF: {filepath}")
            return documents
        
        except ImportError:
            logger.error("PyPDF not installed. Install with: pip install pypdf")
            return []
        except Exception as e:
            logger.error(f"Error loading PDF {filepath}: {str(e)}")
            return []
    
    @staticmethod
    def load_text(filepath: str) -> List[Document]:
        """Load text (.txt, .md) file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            doc = Document(
                page_content=content,
                metadata={
                    "source": filepath,
                    "filename": Path(filepath).name,
                    "type": Path(filepath).suffix.lstrip('.')
                }
            )
            
            logger.info(f"Successfully loaded text file: {filepath}")
            return [doc]
        
        except Exception as e:
            logger.error(f"Error loading text file {filepath}: {str(e)}")
            return []
